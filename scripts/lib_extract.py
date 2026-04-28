"""COA extraction — PDF (Eurofins) + DOCX.

Extracts the envelope metadata (report number, test date, sample name, lot, supersedes)
and a flat list of analyte rows. Downstream (ingest.py) is responsible for:
  - retest collapse (last value wins)
  - superseded/VOID linking
  - unit normalization
  - product-map resolution

This module only pulls raw values — it does not judge them.
"""
from __future__ import annotations

import hashlib
import re
from dataclasses import dataclass, field, asdict
from pathlib import Path
from typing import List, Optional, Tuple

try:
    import pdfplumber  # type: ignore
except ImportError:
    pdfplumber = None

try:
    import docx  # type: ignore  # python-docx
except ImportError:
    docx = None


@dataclass
class AnalyteRow:
    analyte: str
    value_raw: str
    unit_raw: str
    method: Optional[str] = None
    loq: Optional[str] = None
    retest_sequence: int = 0      # 0 = initial, 1 = 1st retest, ...
    source_row_text: str = ""     # audit trail
    panel: str = ""               # heavy_metals | mycotoxins | cga | pesticide | acrylamide | micro | moisture | other


@dataclass
class COAEnvelope:
    report_number: Optional[str] = None
    test_date: Optional[str] = None          # ISO yyyy-mm-dd
    sample_name: Optional[str] = None
    lot_or_po: Optional[str] = None
    supersedes: Optional[str] = None
    lab: Optional[str] = None
    serving_size_g: Optional[float] = None   # Eurofins: "Sample Serving Size 15 g"
    source_file: str = ""
    source_hash: str = ""
    parse_confidence: float = 1.0            # 0..1; <0.6 routes to LOW_CONFIDENCE
    parse_notes: List[str] = field(default_factory=list)
    analytes: List[AnalyteRow] = field(default_factory=list)


# --------------- regex patterns (Eurofins conventions) ---------------

RE_REPORT = re.compile(r"(?:Report\s*(?:No\.?|Number)\s*[:#]?\s*)([0-9]{6,}-?\d*)", re.I)
RE_DATE_STARTED = re.compile(
    r"Date\s*Started\s*[:]?\s*"
    r"([0-9]{1,2}[-/][A-Za-z]{3}[-/][0-9]{2,4}|[0-9]{1,2}[-/][0-9]{1,2}[-/][0-9]{2,4})",
    re.I,
)
RE_SUPERSEDES = re.compile(r"Supe?rcedes?\s*[:]?\s*([0-9]{6,}-?\d*)", re.I)
RE_SAMPLE = re.compile(r"Sample\s*(?:Name|Description|ID)?\s*[:]?\s*(.+)", re.I)
RE_LOT = re.compile(r"(?:Lot|PO|P\.O\.?|Batch)\s*(?:No\.?|#)?\s*[:]?\s*([A-Z0-9-]+)", re.I)
RE_SERVING_SIZE = re.compile(r"(?:Sample\s+)?Serving\s+Size\s+(\d+(?:\.\d+)?)\s*g\b", re.I)

# Eurofins text-line analyte patterns.
# Unit alternation covers the units observed across PROTECT/FLOW/EASE/CALM + component COAs.
_UNIT_ALT = (
    r"mg/Serving\s*Size|mcg/Serving\s*Size|ug/Serving\s*Size|"
    r"g/Serving\s*Size|IU/Serving\s*Size|"
    r"mg/100\s*g|mg/100g|"
    r"mg/kg|mcg/kg|ug/kg|ng/g|ng/kg|"
    r"mg/g|mcg/g|ug/g|"
    r"ppb|ppm|%|"
    r"CFU/g|cfu/g|MPN/g|mpn/g"
)
RE_ANALYTE_LINE = re.compile(
    r"^(?P<name>[A-Za-z0-9][A-Za-z0-9\s\-\(\)\,\./'+]+?)\s+"
    r"(?P<value><\s*[-+]?\d+(?:\.\d+)?|[-+]?\d+(?:\.\d+)?(?:[eE][+-]?\d+)?)\s+"
    r"(?P<unit>" + _UNIT_ALT + r")\s*\*?\s*$",
    re.IGNORECASE,
)
RE_NOT_DETECTED_LINE = re.compile(
    r"^(?P<name>[A-Za-z0-9][A-Za-z0-9\s\-\(\)\,\./'+]+?)\s+"
    r"(?P<verdict>Not\s+Detected(?:\s+at\s+LOQ)?|ND|N\.?D\.?)\s*\*?\s*$",
    re.IGNORECASE,
)
# Dimensionless analytes — water activity, Aw, pH. These have no unit token, just a number.
# Whitelisted so we don't accidentally match arbitrary "<word> <number>" lines.
RE_DIMENSIONLESS_LINE = re.compile(
    r"^(?P<name>"
    r"water\s+activity(?:\s*\(\s*aw\s*\))?"
    r"|aw"
    r"|ph"
    r")\s+"
    r"(?P<value>[-+]?\d+(?:\.\d+)?)\s*\*?\s*$",
    re.IGNORECASE,
)
# Any analyte name must contain at least 3 consecutive letters to count as a real analyte.
RE_HAS_ALPHA = re.compile(r"[A-Za-z]{3,}")
# Text region delimiters — analyte rows live between these markers.
RE_ANALYSIS_START = re.compile(r"^\s*Analysis\s+Result\s*$", re.I | re.M)
RE_METHOD_REF = re.compile(r"^\s*Method\s+References", re.I | re.M)


# --------------- regex patterns (Silliker / Mérieux NutriSciences conventions) ---------------
# These COAs use "COA No: CHG-..." / "COA Date M/D/YY" / "Sample Name: ..." / "P.O.# / ID: ..."
# Analyte table header: "Analyte Result Units Method Reference Test Date[ Loc.]"
# Units carry "(w/w)" suffixes; below-LOQ values are prefixed "<"; "Compounds Not Detected" sections
# list analytes one-per-line with below-LOQ values.

RE_SILLIKER_LAB_SIG = re.compile(
    r"SILLIKER|M[ée]rieux\s*NutriSciences|getresults-core@mxns\.com|COA\s+No\s*:\s*CHG-",
    re.I,
)
RE_SILLIKER_REPORT = re.compile(
    r"(?:COA\s+No|Report\s+No\.?)\s*:?\s*([A-Z]+-\d+-\d+)", re.I,
)
RE_SILLIKER_DATE = re.compile(
    r"(?:COA|Report)\s+Date\s+(\d{1,2}/\d{1,2}/\d{2,4})", re.I,
)
RE_SILLIKER_SUPERSEDES = re.compile(
    r"Supersedes\s*:\s*(None|[A-Z]+-\d+-\d+)", re.I
)
RE_SILLIKER_SAMPLE = re.compile(r"Sample\s*Name\s*:\s*([^\r\n]+)", re.I)
RE_SILLIKER_PO = re.compile(
    r"P\.O\.\s*#\s*/\s*ID\s*:\s*([^\r\n]+?)(?:\s{2,}|\s+Director|\s+Laboratory\s+ID|\s+Location|$)",
    re.I,
)
RE_SILLIKER_DESC1 = re.compile(
    r"Desc\.\s*1\s*:\s*([^\r\n]+?)(?:\s{2,}|\s+Laboratory\s+ID|\s+Condition|$)",
    re.I,
)
RE_SILLIKER_DESC2 = re.compile(
    r"Desc\.\s*2\s*:\s*([^\r\n]+?)(?:\s{2,}|\s+Condition|\s+Temp\s+Rec|$)",
    re.I,
)

# Analyte-region start marker; anchors per page.
RE_SILLIKER_REGION_START = re.compile(
    r"^\s*Analyte\s+Result\s+Units\s+Method\s+Reference\s+Test\s*Date", re.I | re.M
)
# Region end — bottom-of-page disclaimer. Use first N chars of the boilerplate as anchor.
RE_SILLIKER_REGION_END = re.compile(
    r"Results\s+reported\s+herein|_{5,}|\bCertificate\s+of\s+Analysis\b", re.I
)

# Silliker unit alternation. "(w/w)" and "(v/v)" suffixes are optional.
_SILLIKER_UNIT_ALT = (
    r"ppb\s*\(w/w\)|ppm\s*\(w/w\)|%\s*\(w/w\)|"
    r"ppb|ppm|%|"
    r"mcg/kg|ug/kg|mg/kg|ng/g|ng/kg|"
    r"mg/dose|mcg/dose|g/serving|mg/serving|mcg/serving|"
    r"mg\s*a-tocoph/100\s*g|"
    r"/g|MPN/g|CFU/g"
)
RE_SILLIKER_ANALYTE_LINE = re.compile(
    r"^(?P<name>[A-Za-z0-9][A-Za-z0-9\s\-\(\)\,\./'+&]+?)\s+"
    r"(?P<value><\s*[-+]?\d+(?:\.\d+)?|[-+]?\d+(?:\.\d+)?(?:[eE][+-]?\d+)?)\s+"
    r"(?P<unit>" + _SILLIKER_UNIT_ALT + r")"
    # Silliker rows often have trailing method ref + test date + location code.
    # Accept anything (or nothing) after the unit.
    r"(?:\s+.*)?$",
    re.IGNORECASE,
)


def _hash(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1 << 15), b""):
            h.update(chunk)
    return h.hexdigest()[:16]


# --------------- date handling ---------------

_MONTHS = {m: i for i, m in enumerate(
    ["jan", "feb", "mar", "apr", "may", "jun",
     "jul", "aug", "sep", "oct", "nov", "dec"], start=1)}


def iso_date(raw: str) -> Optional[str]:
    if not raw:
        return None
    s = raw.strip().replace("/", "-")
    parts = s.split("-")
    if len(parts) != 3:
        return None
    try:
        if parts[1].isalpha():
            d, mon, y = parts
            m = _MONTHS.get(mon[:3].lower())
            if not m:
                return None
            y = int(y)
            if y < 100:
                y += 2000
            return f"{y:04d}-{m:02d}-{int(d):02d}"
        a, b, c = int(parts[0]), int(parts[1]), int(parts[2])
        if c < 100:
            c += 2000
        # US convention on Eurofins: MM-DD-YYYY
        return f"{c:04d}-{a:02d}-{b:02d}"
    except (ValueError, KeyError):
        return None


# --------------- PDF extraction ---------------

def extract_pdf(path: Path) -> COAEnvelope:
    if pdfplumber is None:
        raise RuntimeError("pdfplumber not installed. pip install pdfplumber")

    env = COAEnvelope(source_file=str(path), source_hash=_hash(path), lab="Eurofins (inferred)")
    text_parts: List[str] = []
    tables: List[List[List[str]]] = []

    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            t = page.extract_text() or ""
            text_parts.append(t)
            for tbl in (page.extract_tables() or []):
                if tbl:
                    tables.append(tbl)

    full_text = "\n".join(text_parts)
    env.parse_notes.append(f"pages={len(text_parts)} tables={len(tables)}")

    # Lab detection — Silliker/Mérieux has a distinct layout and needs a separate parser.
    if RE_SILLIKER_LAB_SIG.search(full_text):
        env.lab = "Silliker / Mérieux NutriSciences"
        _extract_silliker(env, full_text)
        # Confidence heuristic (shared with Eurofins path)
        if not env.report_number:
            env.parse_confidence -= 0.25
            env.parse_notes.append("missing_report_number")
        if not env.test_date:
            env.parse_confidence -= 0.25
            env.parse_notes.append("missing_test_date")
        if not env.analytes:
            env.parse_confidence -= 0.4
            env.parse_notes.append("no_analyte_rows_extracted")
        env.parse_confidence = max(0.0, min(1.0, env.parse_confidence))
        return env

    if m := RE_REPORT.search(full_text):
        env.report_number = m.group(1)
    if m := RE_DATE_STARTED.search(full_text):
        env.test_date = iso_date(m.group(1))
    if m := RE_SUPERSEDES.search(full_text):
        env.supersedes = m.group(1)
    if m := RE_SAMPLE.search(full_text):
        env.sample_name = m.group(1).strip().splitlines()[0][:200]
    if m := RE_LOT.search(full_text):
        env.lot_or_po = m.group(1)
    if m := RE_SERVING_SIZE.search(full_text):
        try:
            env.serving_size_g = float(m.group(1))
        except ValueError:
            pass

    # Analyte rows from tables — generic row scanner.
    for tbl in tables:
        rows = [[(c or "").strip() for c in row] for row in tbl if row]
        if not rows:
            continue
        header = [c.lower() for c in rows[0]]
        # Heuristic: look for columns containing analyte/result/unit
        def col(*names):
            for i, h in enumerate(header):
                if any(n in h for n in names):
                    return i
            return None

        c_analyte = col("analyte", "parameter", "test", "compound")
        c_result  = col("result", "value", "amount", "conc")
        c_unit    = col("unit", "uom")
        c_method  = col("method", "technique")
        c_loq     = col("loq", "lod", "limit")

        if c_analyte is None or c_result is None:
            continue

        for r in rows[1:]:
            if len(r) <= max(c_analyte, c_result):
                continue
            name = r[c_analyte]
            val  = r[c_result]
            if not name or not val:
                continue
            row = AnalyteRow(
                analyte=name,
                value_raw=val,
                unit_raw=r[c_unit] if c_unit is not None and c_unit < len(r) else "",
                method=r[c_method] if c_method is not None and c_method < len(r) else None,
                loq=r[c_loq] if c_loq is not None and c_loq < len(r) else None,
                source_row_text=" | ".join(r),
                panel=_guess_panel(name),
            )
            env.analytes.append(row)

    # Text-line fallback — Eurofins COAs put results in free-form text, not tables.
    # Always run this; union with any table-derived rows and dedupe later.
    text_rows = _extract_analytes_from_text(full_text)
    if text_rows:
        # Prefer text rows when tables yielded nothing; otherwise append
        # only analytes we didn't already see from tables.
        seen_names = {a.analyte.strip().lower() for a in env.analytes}
        for r in text_rows:
            if r.analyte.strip().lower() not in seen_names:
                env.analytes.append(r)
        env.parse_notes.append(f"text_rows={len(text_rows)}")

    # Confidence heuristic
    if not env.report_number:
        env.parse_confidence -= 0.25
        env.parse_notes.append("missing_report_number")
    if not env.test_date:
        env.parse_confidence -= 0.25
        env.parse_notes.append("missing_test_date")
    if not env.analytes:
        env.parse_confidence -= 0.4
        env.parse_notes.append("no_analyte_rows_extracted")
    env.parse_confidence = max(0.0, min(1.0, env.parse_confidence))
    return env


def _extract_silliker(env: COAEnvelope, full_text: str) -> None:
    """Silliker/Mérieux extractor.

    Populates env.report_number, env.test_date, env.supersedes, env.sample_name,
    env.lot_or_po, and env.analytes.
    """
    if m := RE_SILLIKER_REPORT.search(full_text):
        env.report_number = m.group(1)
    if m := RE_SILLIKER_DATE.search(full_text):
        env.test_date = iso_date(m.group(1))
    if m := RE_SILLIKER_SUPERSEDES.search(full_text):
        sup = m.group(1).strip()
        env.supersedes = None if sup.lower() == "none" else sup
    # Sample name: prefer "Sample Name:" then "Desc. 1:" then "Desc. 2:" (older COAs used Desc.)
    if m := RE_SILLIKER_SAMPLE.search(full_text):
        env.sample_name = m.group(1).strip()[:200]
    elif m := RE_SILLIKER_DESC1.search(full_text):
        env.sample_name = m.group(1).strip()[:200]
        # If Desc. 2 is present, use it as additional context in lot_or_po
        if d2 := RE_SILLIKER_DESC2.search(full_text):
            extra = d2.group(1).strip()
            if extra and extra.lower() not in env.sample_name.lower():
                env.sample_name = f"{env.sample_name} / {extra}"[:200]
    elif m := RE_SILLIKER_DESC2.search(full_text):
        env.sample_name = m.group(1).strip()[:200]
    # Lot / PO: prefer explicit P.O.# field
    if m := RE_SILLIKER_PO.search(full_text):
        env.lot_or_po = m.group(1).strip()[:200]

    # Analyte regions — one per "Analyte Result Units Method Reference..." marker.
    rows: List[AnalyteRow] = []
    seen_names = set()
    # Find each region start; slice to the next disclaimer (or next header) for the analyte body.
    for start_m in RE_SILLIKER_REGION_START.finditer(full_text):
        region_start = start_m.end()
        end_m = RE_SILLIKER_REGION_END.search(full_text, region_start)
        region = full_text[region_start : end_m.start() if end_m else len(full_text)]

        for raw_line in region.splitlines():
            line = raw_line.strip()
            if not line:
                continue
            low = line.lower()
            # Skip method-reference sub-headers (lines with a date but no value+unit pattern)
            # e.g. "Aflatoxin by HPLC AOAC 991.31 (Mod.) 3/27/18"
            # These won't match RE_SILLIKER_ANALYTE_LINE so just try the match.
            # Skip common section captions / meta lines
            if low.startswith((
                "compounds detected", "compounds not detected",
                "(none detected)",
                "laboratory id", "condition rec", "temp rec",
                "sample name", "additional field",
                "desc. 1", "desc. 2", "p.o.#",
                "noted test locations",
                "analytical results",
                "page ", "total chloroge",  # keep total chlorogenic via match below
            )):
                # "total chloroge" is a defensive prefix — the actual Total Chlorogenic Acid
                # line has value+unit and will match below before we hit this.
                if low.startswith("total chloroge"):
                    pass  # let it fall through to matcher
                else:
                    continue
            m = RE_SILLIKER_ANALYTE_LINE.match(line)
            if not m:
                continue
            name = m.group("name").strip()
            # Reject header remnants
            if len(name) < 2 or not RE_HAS_ALPHA.search(name):
                continue
            low_name = name.lower()
            # Drop noise names that sometimes slip through (page footer fragments, etc.)
            if low_name in {"page", "result", "analyte", "units", "result units"}:
                continue
            # De-dup within the file by lowercase name (retest_collapse runs later in pipeline)
            if low_name in seen_names:
                # Allow multi-region duplicates only if this is a retest — keep as a second row
                # so collapse_retests can apply "last wins" logic.
                pass
            seen_names.add(low_name)
            value = m.group("value").strip().replace(" ", "")
            unit = m.group("unit").strip()
            rows.append(AnalyteRow(
                analyte=name,
                value_raw=value,
                unit_raw=unit,
                source_row_text=line,
                panel=_guess_panel(name),
            ))

    env.analytes.extend(rows)
    if rows:
        env.parse_notes.append(f"silliker_rows={len(rows)}")


def _extract_analytes_from_text(full_text: str) -> List[AnalyteRow]:
    """Parse Eurofins-style free-text analyte lines.

    Region of interest: between 'Analysis Result' and 'Method References'.
    Matches two shapes:
      1. '<name> <value> <unit>'  e.g. 'Lead <5.00 ppb', 'Chlorogenic Acid 103 mg/Serving Size'
      2. '<name> Not Detected at LOQ'
    Skips section headers (no value/unit) and metadata rows like 'Project ID ...'.
    """
    rows: List[AnalyteRow] = []
    # Narrow to the result region if markers are present; otherwise scan whole text.
    start_m = RE_ANALYSIS_START.search(full_text)
    end_m = RE_METHOD_REF.search(full_text, start_m.end() if start_m else 0)
    region = full_text[start_m.end() if start_m else 0 : end_m.start() if end_m else len(full_text)]

    # Skip lines that are clearly metadata (defensive — the region slice should already exclude them).
    METADATA_PREFIXES = (
        "report number", "report date", "report status", "project id",
        "po number", "lot number", "sample serving size", "date started",
        "sampled", "online order", "printed:", "certificate of analysis",
        "description", "login date", "receipt date", "receipt condition",
        "sample name", "sample results",
    )

    for raw_line in region.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        low = line.lower()
        if any(low.startswith(p) for p in METADATA_PREFIXES):
            continue

        m = RE_ANALYTE_LINE.match(line)
        if m:
            name = m.group("name").strip()
            # Reject if "name" is actually a unit/method fragment
            if len(name) < 2 or name.lower() in {"result", "analyte"}:
                continue
            # Require at least 3 consecutive letters — rejects pure numeric/code fragments
            if not RE_HAS_ALPHA.search(name):
                continue
            value = m.group("value").strip().replace(" ", "")
            unit = m.group("unit").strip()
            rows.append(AnalyteRow(
                analyte=name,
                value_raw=value,
                unit_raw=unit,
                source_row_text=line,
                panel=_guess_panel(name),
            ))
            continue

        nd = RE_NOT_DETECTED_LINE.match(line)
        if nd:
            name = nd.group("name").strip()
            if len(name) < 2:
                continue
            # Skip if it looks like a section header fragment
            if name.lower().startswith(("screened", "method", "testing")):
                # But keep 'Screened pesticides Not Detected at LOQ' since it conveys info
                if "screened" in name.lower() and "pesticide" in name.lower():
                    pass
                else:
                    continue
            rows.append(AnalyteRow(
                analyte=name,
                value_raw="Not Detected",
                unit_raw="",
                source_row_text=line,
                panel=_guess_panel(name),
            ))
            continue

        # Dimensionless-analyte fallback (water activity, Aw, pH).
        dim = RE_DIMENSIONLESS_LINE.match(line)
        if dim:
            name = dim.group("name").strip()
            value = dim.group("value").strip()
            rows.append(AnalyteRow(
                analyte=name,
                value_raw=value,
                unit_raw="",
                source_row_text=line,
                panel=_guess_panel(name),
            ))

    return rows


# --------------- DOCX extraction ---------------

def extract_docx(path: Path) -> COAEnvelope:
    if docx is None:
        raise RuntimeError("python-docx not installed. pip install python-docx")

    env = COAEnvelope(source_file=str(path), source_hash=_hash(path), lab="Eurofins (inferred)")
    doc = docx.Document(str(path))

    paragraphs = "\n".join(p.text for p in doc.paragraphs)
    if m := RE_REPORT.search(paragraphs):
        env.report_number = m.group(1)
    if m := RE_DATE_STARTED.search(paragraphs):
        env.test_date = iso_date(m.group(1))
    if m := RE_SUPERSEDES.search(paragraphs):
        env.supersedes = m.group(1)
    if m := RE_SAMPLE.search(paragraphs):
        env.sample_name = m.group(1).strip().splitlines()[0][:200]
    if m := RE_LOT.search(paragraphs):
        env.lot_or_po = m.group(1)

    for tbl in doc.tables:
        rows_raw = [[cell.text.strip() for cell in row.cells] for row in tbl.rows]
        if not rows_raw or len(rows_raw[0]) < 2:
            continue
        header = [c.lower() for c in rows_raw[0]]
        try:
            c_analyte = next(i for i, h in enumerate(header)
                             if any(k in h for k in ("analyte", "parameter", "test", "compound")))
            c_result  = next(i for i, h in enumerate(header)
                             if any(k in h for k in ("result", "value", "amount")))
        except StopIteration:
            env.parse_notes.append(f"docx_table_header_unrecognized: {header[:4]}")
            env.parse_confidence -= 0.1
            continue

        for r in rows_raw[1:]:
            if len(r) <= max(c_analyte, c_result):
                continue
            name, val = r[c_analyte], r[c_result]
            if not name or not val:
                continue
            unit = ""
            for i, h in enumerate(header):
                if i < len(r) and ("unit" in h or "uom" in h):
                    unit = r[i]
                    break
            env.analytes.append(AnalyteRow(
                analyte=name, value_raw=val, unit_raw=unit,
                source_row_text=" | ".join(r), panel=_guess_panel(name),
            ))

    if not env.analytes:
        env.parse_confidence -= 0.4
        env.parse_notes.append("no_analyte_rows_extracted")
    env.parse_confidence = max(0.0, min(1.0, env.parse_confidence))
    return env


def _guess_panel(name: str) -> str:
    n = name.lower()
    if any(k in n for k in ("lead", "cadmium", "arsenic", "mercury")):
        return "heavy_metals"
    if any(k in n for k in ("ochratoxin", "aflatoxin")):
        return "mycotoxins"
    # Eurofins has spelling typos ("Caffeolyquinic", "Caffeoylquiic"); match any plausible variant.
    if ("chlorogenic" in n or n.endswith("cqa")
            or "caffeoylquinic" in n or "caffeolyquinic" in n
            or "caffeoylquiic" in n or "caffeolyquiic" in n
            or "dicaffeoylquinic" in n or "dicaffeolyquinic" in n
            or "dicaffeoylquiic" in n):
        return "cga"
    if "acrylamide" in n:
        return "acrylamide"
    if "yeast" in n or "mold" in n:
        return "micro"
    if "moisture" in n or "water activity" in n or n.strip() == "aw":
        return "moisture"
    if "pesticide" in n or "gc-ms" in n or "lc-ms" in n or "glyphosate" in n or "ampa" in n:
        return "pesticide"
    if "gluten" in n:
        return "allergen"
    return "other"


def extract(path: Path) -> COAEnvelope:
    ext = path.suffix.lower()
    if ext == ".pdf":
        return extract_pdf(path)
    if ext == ".docx":
        return extract_docx(path)
    raise ValueError(f"Unsupported file type: {ext}")


# --------------- retest collapse ---------------

def collapse_retests(env: COAEnvelope) -> COAEnvelope:
    """For rows with the same analyte name, the last occurrence wins.
    Earlier values are preserved in a retest_history list on the kept row.

    Eurofins convention: retests appear as separate rows in sequence in the
    same table. This function is order-preserving — it assumes parsing kept
    table order intact.
    """
    seen: dict[str, AnalyteRow] = {}
    history: dict[str, list] = {}
    final: List[AnalyteRow] = []

    for row in env.analytes:
        key = row.analyte.strip().lower()
        if key in seen:
            history.setdefault(key, []).append({
                "value_raw": seen[key].value_raw,
                "unit_raw":  seen[key].unit_raw,
                "source_row_text": seen[key].source_row_text,
            })
            row.retest_sequence = seen[key].retest_sequence + 1
        seen[key] = row

    # rebuild final list in original first-seen order
    order = []
    order_seen = set()
    for row in env.analytes:
        key = row.analyte.strip().lower()
        if key not in order_seen:
            order.append(key)
            order_seen.add(key)

    for key in order:
        row = seen[key]
        row.source_row_text = row.source_row_text
        if key in history:
            # attach retest history as parse note on envelope so JSON stays audit-clean
            env.parse_notes.append(
                f"retest:{row.analyte}:kept='{row.value_raw} {row.unit_raw}' prior={history[key]}"
            )
        final.append(row)
    env.analytes = final
    return env
